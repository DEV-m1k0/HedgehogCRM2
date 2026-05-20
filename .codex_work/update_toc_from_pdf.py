from __future__ import annotations

import sys
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from pypdf import PdfReader

ROOT = Path("/Users/antonknyazev/Desktop/HedgehogCRM2")
DOCX = ROOT / "diplome_final.docx"
PDF = Path(os.environ.get("PDF_PATH", "/tmp/hedgehog_table_fix_render/diplome_final.pdf"))

sys.path.insert(0, str(ROOT / ".codex_work"))
from build_diplome_final import TOC_ITEMS, patch_ooxml_text, update_toc  # noqa: E402


def norm(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split()).lower()


def toc_search_text(label: str, key: str) -> str:
    if label.startswith("Приложение "):
        return norm(label.replace(".", "", 1))
    return norm(key)


def find_pages(pdf_path: Path, skip_pages: int = 1) -> dict[str, int]:
    reader = PdfReader(str(pdf_path))
    page_texts = [(idx, norm(page.extract_text() or "")) for idx, page in enumerate(reader.pages, start=1)]
    pages: dict[str, int] = {}
    for label, key in TOC_ITEMS:
        needle = toc_search_text(label, key)
        for page_no, text in page_texts:
            if page_no <= skip_pages:
                continue
            if needle in text:
                pages[key] = page_no
                break
    print(f"pdf_pages={len(reader.pages)}")
    for _, key in TOC_ITEMS:
        print(f"{key}: {pages.get(key)}")
    return pages


def force_toc_split(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith("Список литературы"):
            p.add_run().add_break(WD_BREAK.PAGE)
            return


def main() -> None:
    skip_pages = int(os.environ.get("TOC_SKIP_PAGES", "1"))
    pages = find_pages(PDF, skip_pages=skip_pages)
    doc = Document(DOCX)
    update_toc(doc, pages)
    force_toc_split(doc)
    doc.save(DOCX)
    patch_ooxml_text(DOCX, page_count=len(PdfReader(str(PDF)).pages))
    print(f"updated {DOCX}")


if __name__ == "__main__":
    main()
