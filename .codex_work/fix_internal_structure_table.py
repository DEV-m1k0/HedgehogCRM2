from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path("/Users/antonknyazev/Desktop/HedgehogCRM2/diplome_final.docx")


TABLE1_ROWS = [
    (
        "Наименование объекта",
        "Назначение",
        "Свойства и события",
    ),
    (
        "models.py",
        "Файл описания ORM-моделей данных приложения.",
        "Содержит классы Role, User, Client, Course, StudyGroup, Lesson, Attendance, Deal, Task, UserSession и AuditLog. Основные свойства задаются через mapped_column; связи между сущностями описаны через relationship. Для User предусмотрены методы set_password и check_password.",
    ),
    (
        "main.py",
        "Основной файл серверного приложения FastAPI.",
        "Создаёт экземпляр приложения, подключает CORS, статические медиафайлы, маршрутизаторы и middleware журналирования. В событии startup создаёт таблицы, проверяет актуальность схемы, заполняет роли и запускает сервис напоминаний.",
    ),
    (
        "security.py",
        "Модуль безопасности и авторизации.",
        "Содержит параметры JWT, функции создания access/refresh-токенов, декодирования токена, хеширования refresh-токена и проверки пароля пользователя.",
    ),
    (
        "db.py",
        "Модуль подключения к базе данных.",
        "Использует DATABASE_URL, создаёт SQLAlchemy engine и фабрику SessionLocal. Для SQLite задаёт check_same_thread, для PostgreSQL использует стандартное подключение SQLAlchemy.",
    ),
    (
        "deps.py",
        "Модуль зависимостей FastAPI.",
        "Предоставляет get_db, get_current_user и проверки ролей. Обрабатывает события отсутствующей, завершённой или недействительной пользовательской сессии.",
    ),
    (
        "router.py",
        "Агрегатор маршрутов REST API.",
        "Создаёт общий APIRouter и подключает маршрутизаторы auth, clients, courses, groups, schedule, deals, tasks, archive, admin и meta с соответствующими префиксами.",
    ),
    (
        "auth.py",
        "Маршрутизатор регистрации, входа и сессий.",
        "Обрабатывает регистрацию, авторизацию, обновление токенов и выход из системы. При входе создаёт пользовательскую сессию и refresh-token cookie.",
    ),
    (
        "clients.py",
        "Маршрутизатор карточек учеников и статистики.",
        "Реализует создание, просмотр, изменение и архивирование учеников. Для карточки ученика рассчитывает статистику посещаемости и сведения, используемые AI-модулем.",
    ),
    (
        "schedule.py",
        "Маршрутизатор расписания, занятий и посещаемости.",
        "Управляет одиночными и повторяющимися занятиями, отменой, архивированием, журналом посещаемости, отработками и статусом проведения занятия.",
    ),
    (
        "deals.py и tasks.py",
        "Маршрутизаторы сделок и задач сотрудников.",
        "Deals.py ведёт воронку продаж, суммы, этапы и статусы сделок. Tasks.py создаёт задачи, назначает исполнителя, приоритет, срок и статус выполнения.",
    ),
    (
        "archive.py и admin.py",
        "Маршрутизаторы архива и административного контроля.",
        "Archive.py восстанавливает архивированные сущности. Admin.py предоставляет сведения о сотрудниках, активных сессиях, журнале аудита и действия администратора по управлению пользователями.",
    ),
]


def set_run_font(run, size: float = 10, name: str = "Times New Roman", bold: bool | None = None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_cell_margins(cell, margin: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    # Methodology: 14 pt text, single line spacing for tables. 10 pt is only for full listings.
    set_run_font(run, 14, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, 100)


def set_table_width(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    widths = [Cm(3.1), Cm(5.2), Cm(8.7)]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(Cm(17).twips)))
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(widths[idx].twips)))


def normalize_toc(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if "\t" in text and (text.startswith("Введение") or text.startswith("1 ") or text.startswith("Приложение")):
            p.paragraph_format.tab_stops.clear_all()
            p.paragraph_format.tab_stops.add_tab_stop(Cm(16.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def main() -> None:
    doc = Document(DOCX)
    table = doc.tables[0]

    # Rebuild table rows while preserving the existing table object location.
    for tr in list(table._tbl.tr_lst):
        table._tbl.remove(tr)
    for row_idx, values in enumerate(TABLE1_ROWS):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            set_cell_text(cells[col_idx], value, bold=(row_idx == 0))
    set_table_width(table)
    normalize_toc(doc)
    doc.save(DOCX)
    print(f"fixed {DOCX}")


if __name__ == "__main__":
    main()
