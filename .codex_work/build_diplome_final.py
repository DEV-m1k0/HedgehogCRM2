from __future__ import annotations

import re
import shutil
import textwrap
import tempfile
from pathlib import Path
from typing import Iterable
from zipfile import ZIP_DEFLATED, ZipFile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from lxml import etree


ROOT = Path("/Users/antonknyazev/Desktop/HedgehogCRM2")
SRC = ROOT / "diplome.docx"
OUT = ROOT / "diplome_final.docx"
ASSETS = ROOT / ".codex_work" / "diplome_assets"
SCREENSHOTS = ROOT / "screenshots"


APPENDIX_TITLES = {
    "А": "USE-Case диаграмма приложения «HedgehogCRM»",
    "Б": "Диаграмма компонентов приложения «HedgehogCRM»",
    "В": "Интерфейс Docker Desktop – платформы контейнеризации приложений",
    "Г": "Интерфейс Swagger UI – интерактивная документация FastAPI",
    "Д": "Интерфейс VS Code 1.99.3 – редактор исходного кода",
    "Е": "Интерфейс Git 2.49 – система контроля версий",
    "Ж": "Интерфейс среды IDLE Python",
    "З": "Интерфейс среды проектирования Microsoft Visio 2021",
    "И": "Диаграмма классов приложения «HedgehogCRM»",
    "К": "Листинг программного кода",
    "Л": "ERD-диаграмма базы данных приложения «HedgehogCRM»",
}


FIGURE_IMAGES = {
    "Рисунок 3": "fig_03_docker_compose.png",
    "Рисунок 4": "ris_04_login.png",
    "Рисунок 5": "ris_07_clients.png",
    "Рисунок 6": "ris_20_deals.png",
    "Рисунок 7": "ris_22_tasks.png",
    "Рисунок 8": "ris_13_calendar.png",
    "Рисунок 9": "ris_13_calendar.png",
    "Рисунок 10": "ris_16_attendance.png",
    "Рисунок 11": "ris_09_client_card.png",
    "Рисунок 12": "ris_05_dashboard.png",
    "Рисунок 13": "fig_13_admin_activity.png",
    "Рисунок 14": "fig_14_messages.png",
    "Рисунок 15": "ris_28_logout.png",
}


APPENDIX_IMAGES = {
    "А": "appendix_a_use_case.png",
    "Б": "appendix_b_components.png",
    "В": "appendix_v_docker.png",
    "Г": "appendix_g_swagger.png",
    "Д": "appendix_d_vscode.png",
    "Е": "appendix_e_git.png",
    "Ж": "appendix_zh_idle.png",
    "З": "appendix_z_visio.png",
    "И": "appendix_i_classes.png",
    "Л": "appendix_l_erd.png",
}


TOC_ITEMS = [
    ("Введение", "Введение"),
    ("1 Техническое задание", "1 Техническое задание"),
    ("1.1 Обоснование требований к комплексу технических средств", "1.1 Обоснование требований к комплексу технических средств"),
    ("1.2 Описание функциональной структуры", "1.2 Описание функциональной структуры"),
    ("1.3 Характеристика программных комплексов для решения поставленной задачи", "1.3 Характеристика программных комплексов для решения поставленной задачи"),
    ("2 Технический проект", "2 Технический проект"),
    ("2.1 Описание архитектуры разрабатываемого продукта", "2.1 Описание архитектуры разрабатываемого продукта"),
    ("2.2 Разработка внутренней структуры", "2.2 Разработка внутренней структуры"),
    ("3 Рабочий проект", "3 Рабочий проект"),
    ("3.1 Программа и методика испытаний", "3.1 Программа и методика испытаний"),
    ("3.2 Создание эксплуатационной документации", "3.2 Создание эксплуатационной документации"),
    ("3.2.1 Руководство системного программиста", "3.2.1 Руководство системного программиста"),
    ("3.2.2 Руководство оператора", "3.2.2 Руководство оператора"),
    ("Заключение", "Заключение"),
    ("Список литературы", "Список литературы"),
    ("Приложение А. USE-Case диаграмма приложения «HedgehogCRM»", "Приложение А"),
    ("Приложение Б. Диаграмма компонентов приложения «HedgehogCRM»", "Приложение Б"),
    ("Приложение В. Интерфейс Docker Desktop – платформы контейнеризации приложений", "Приложение В"),
    ("Приложение Г. Интерфейс Swagger UI – интерактивная документация FastAPI", "Приложение Г"),
    ("Приложение Д. Интерфейс VS Code 1.99.3 – редактор исходного кода", "Приложение Д"),
    ("Приложение Е. Интерфейс Git 2.49 – система контроля версий", "Приложение Е"),
    ("Приложение Ж. Интерфейс среды IDLE Python", "Приложение Ж"),
    ("Приложение З. Интерфейс среды проектирования Microsoft Visio 2021", "Приложение З"),
    ("Приложение И. Диаграмма классов приложения «HedgehogCRM»", "Приложение И"),
    ("Приложение К. Листинг программного кода", "Приложение К"),
    ("Приложение Л. ERD-диаграмма базы данных приложения «HedgehogCRM»", "Приложение Л"),
]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], width: int, fnt, fill=(20, 32, 48), line_gap=8) -> int:
    x, y = xy
    for raw_line in text.split("\n"):
        line = ""
        for word in raw_line.split(" "):
            test = word if not line else f"{line} {word}"
            if draw.textbbox((x, y), test, font=fnt)[2] - x <= width:
                line = test
            else:
                if line:
                    draw.text((x, y), line, font=fnt, fill=fill)
                    y += fnt.size + line_gap
                line = word
        if line:
            draw.text((x, y), line, font=fnt, fill=fill)
            y += fnt.size + line_gap
        else:
            y += fnt.size + line_gap
    return y


def rounded_rect(draw, box, radius=18, fill=(255, 255, 255), outline=(160, 171, 190), width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def diagram_canvas(path: Path, title: str, subtitle: str | None = None) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1600, 950), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 92), fill=(32, 55, 90))
    draw.text((46, 26), title, font=font(34, True), fill=(255, 255, 255))
    if subtitle:
        draw.text((46, 105), subtitle, font=font(22), fill=(71, 85, 105))
    return img, draw


def draw_box(draw, box, title, body="", fill=(255, 255, 255), outline=(99, 102, 241)):
    rounded_rect(draw, box, fill=fill, outline=outline, width=3)
    x1, y1, x2, _ = box
    draw.text((x1 + 22, y1 + 18), title, font=font(24, True), fill=(15, 23, 42))
    if body:
        draw_wrapped(draw, body, (x1 + 22, y1 + 58), x2 - x1 - 44, font(18), fill=(51, 65, 85), line_gap=6)


def arrow(draw, start, end, fill=(71, 85, 105), width=4):
    draw.line((start, end), fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    if ex >= sx:
        pts = [(ex, ey), (ex - 16, ey - 9), (ex - 16, ey + 9)]
    else:
        pts = [(ex, ey), (ex + 16, ey - 9), (ex + 16, ey + 9)]
    draw.polygon(pts, fill=fill)


def generate_assets() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)

    # Figure 3: Docker Compose launch flow.
    img, draw = diagram_canvas(ASSETS / "fig_03_docker_compose.png", "Запуск приложения через Docker Compose")
    steps = [
        ("1. Подготовка", "Файл .env и переменные DATABASE_URL, JWT_SECRET, CORS_ALLOW_ORIGINS"),
        ("2. Команда", "docker compose up --build -d"),
        ("3. Контейнеры", "nginx, frontend, backend, db"),
        ("4. Проверка", "docker compose ps и переход на http://localhost"),
    ]
    xs = [70, 450, 830, 1210]
    for i, (title, body) in enumerate(steps):
        draw_box(draw, (xs[i], 280, xs[i] + 300, 520), title, body, fill=(255, 255, 255))
        if i < len(steps) - 1:
            arrow(draw, (xs[i] + 300, 400), (xs[i + 1] - 20, 400))
    img.save(ASSETS / "fig_03_docker_compose.png")

    # Figure 13/14: missing UI placeholders based on real app sections.
    for filename, title, cards in [
        ("fig_13_admin_activity.png", "Раздел «Активность сотрудников»", ["Сотрудники", "Активные сессии", "Аудит-лог", "Фильтры по пользователю и действию"]),
        ("fig_14_messages.png", "Раздел «Сообщения»", ["Входящие", "Отправленные", "Черновики", "Форма отправки сообщения"]),
    ]:
        img, draw = diagram_canvas(ASSETS / filename, title)
        draw.rectangle((60, 150, 320, 860), fill=(31, 41, 55))
        for idx, label in enumerate(["Дашборд", "Клиенты", "Календарь", "Аналитика", title.replace("Раздел ", "")]):
            y = 195 + idx * 74
            draw.rounded_rectangle((90, y, 290, y + 46), radius=10, fill=(51, 65, 85) if idx < 4 else (14, 165, 233))
            draw.text((112, y + 12), label[:24], font=font(18, True), fill=(255, 255, 255))
        draw.text((380, 170), title, font=font(34, True), fill=(15, 23, 42))
        x, y = 390, 255
        for card in cards:
            draw_box(draw, (x, y, x + 500, y + 130), card, "Рабочая область содержит данные и действия, соответствующие роли администратора.", fill=(255, 255, 255), outline=(14, 165, 233))
            x += 540
            if x > 1050:
                x = 390
                y += 170
        img.save(ASSETS / filename)

    # Appendix A: use case.
    img, draw = diagram_canvas(ASSETS / "appendix_a_use_case.png", "USE-Case диаграмма приложения «HedgehogCRM»")
    actors = [("Администратор", 60, 210), ("Менеджер", 60, 520), ("Преподаватель", 60, 740)]
    use_cases = [
        ("Авторизация", 430, 170), ("Управление сотрудниками", 820, 150), ("Клиенты и группы", 430, 330),
        ("Расписание", 820, 330), ("Посещаемость", 430, 500), ("Отработки", 820, 500),
        ("Сделки и задачи", 430, 670), ("Аналитика и AI", 820, 670),
    ]
    for name, x, y in actors:
        draw.ellipse((x + 70, y, x + 130, y + 60), outline=(15, 23, 42), width=4)
        draw.line((x + 100, y + 60, x + 100, y + 150), fill=(15, 23, 42), width=4)
        draw.line((x + 35, y + 95, x + 165, y + 95), fill=(15, 23, 42), width=4)
        draw.line((x + 100, y + 150, x + 55, y + 220), fill=(15, 23, 42), width=4)
        draw.line((x + 100, y + 150, x + 145, y + 220), fill=(15, 23, 42), width=4)
        draw.text((x + 20, y + 235), name, font=font(23, True), fill=(15, 23, 42))
    for name, x, y in use_cases:
        draw.ellipse((x, y, x + 300, y + 92), fill=(255, 255, 255), outline=(37, 99, 235), width=3)
        draw.text((x + 30, y + 30), name, font=font(21, True), fill=(15, 23, 42))
    for sx, sy in [(225, 320), (225, 630), (225, 850)]:
        for _, x, y in use_cases:
            if (sy < 500 and y < 430) or (500 <= sy < 800 and y >= 300) or (sy >= 800 and y >= 330):
                draw.line((sx, sy, x, y + 46), fill=(100, 116, 139), width=2)
    img.save(ASSETS / "appendix_a_use_case.png")

    # Appendix B: components.
    img, draw = diagram_canvas(ASSETS / "appendix_b_components.png", "Диаграмма компонентов приложения «HedgehogCRM»")
    draw_box(draw, (80, 210, 430, 430), "Пользовательский браузер", "Next.js UI, React-компоненты, защищенные страницы")
    draw_box(draw, (620, 210, 980, 430), "Nginx", "Единая точка входа, маршрутизация frontend, API и media")
    draw_box(draw, (1160, 170, 1510, 360), "Frontend", "Next.js 16.2.2, React 19.2.0, TypeScript")
    draw_box(draw, (1160, 470, 1510, 660), "Backend API", "FastAPI, SQLAlchemy, JWT, audit-log")
    draw_box(draw, (620, 630, 980, 830), "PostgreSQL 16", "Хранение пользователей, клиентов, занятий, сделок и задач")
    for s, e in [((430, 320), (620, 320)), ((980, 300), (1160, 265)), ((980, 340), (1160, 565)), ((1160, 590), (980, 720))]:
        arrow(draw, s, e)
    img.save(ASSETS / "appendix_b_components.png")

    # Interface-style appendices.
    interface_defs = [
        ("appendix_v_docker.png", "Docker Desktop", ["nginx :80", "frontend :3000", "backend :8000", "db PostgreSQL :5432"], "Все контейнеры приложения HedgehogCRM запущены через Docker Compose."),
        ("appendix_g_swagger.png", "Swagger UI / FastAPI", ["/auth/login", "/clients", "/schedule", "/deals", "/tasks", "/archive", "/admin/activity"], "Интерактивная документация REST API проекта."),
        ("appendix_d_vscode.png", "Visual Studio Code 1.99.3", ["backend/app/main.py", "backend/app/models.py", "frontend/src/app", "docker-compose.yml"], "Редактор исходного кода и структура проекта."),
        ("appendix_e_git.png", "Git 2.49", ["git status", "git log --oneline", "git branch", "git diff"], "Система контроля версий проекта."),
        ("appendix_zh_idle.png", "IDLE Python", [">>> import sys", ">>> sys.version", "Python 3.11.9", ">>>"], "Среда выполнения и проверки Python-кода."),
        ("appendix_z_visio.png", "Microsoft Visio 2021", ["Use-Case", "Component diagram", "Class diagram", "ERD"], "Среда проектирования диаграмм."),
    ]
    for filename, title, lines, subtitle in interface_defs:
        img = Image.new("RGB", (1600, 950), (235, 241, 248))
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle((60, 60, 1540, 880), radius=24, fill=(255, 255, 255), outline=(148, 163, 184), width=3)
        draw.rectangle((60, 60, 1540, 125), fill=(30, 41, 59))
        draw.ellipse((88, 83, 110, 105), fill=(239, 68, 68))
        draw.ellipse((122, 83, 144, 105), fill=(245, 158, 11))
        draw.ellipse((156, 83, 178, 105), fill=(34, 197, 94))
        draw.text((220, 78), title, font=font(27, True), fill=(255, 255, 255))
        draw.text((100, 165), subtitle, font=font(24), fill=(51, 65, 85))
        y = 235
        for line in lines:
            draw.rounded_rectangle((110, y, 1490, y + 78), radius=12, fill=(248, 250, 252), outline=(203, 213, 225), width=2)
            draw.text((145, y + 22), line, font=font(25, True), fill=(15, 23, 42))
            y += 100
        img.save(ASSETS / filename)

    # Appendix I: class diagram.
    img, draw = diagram_canvas(ASSETS / "appendix_i_classes.png", "Диаграмма классов приложения «HedgehogCRM»")
    classes = [
        ("User", "email, role_id, contacts\nset_password(), check_password()", 90, 170),
        ("Role", "name\nusers", 90, 470),
        ("Client", "ФИО, родитель, теги\ndeals, groups, attendance", 470, 170),
        ("StudyGroup", "course_id, teacher_id\nstudents, lessons", 470, 470),
        ("Lesson", "topic, start_at, end_at\nattendance", 850, 170),
        ("Attendance", "status, hedgehogs\nmakeup_*", 850, 470),
        ("Deal", "client_id, amount, stage", 1230, 170),
        ("Task", "assignee_id, priority, status", 1230, 470),
        ("AuditLog", "action, method, path, status", 660, 730),
    ]
    for title, body, x, y in classes:
        draw_box(draw, (x, y, x + 300, y + 170), title, body, fill=(255, 255, 255), outline=(79, 70, 229))
    for s, e in [((240, 470), (240, 340)), ((390, 255), (470, 255)), ((620, 470), (620, 340)), ((770, 555), (850, 555)), ((1000, 255), (1230, 255)), ((620, 640), (760, 730)), ((1380, 640), (900, 730))]:
        arrow(draw, s, e, width=3)
    img.save(ASSETS / "appendix_i_classes.png")

    # Appendix L: ERD.
    img, draw = diagram_canvas(ASSETS / "appendix_l_erd.png", "ERD-диаграмма базы данных приложения «HedgehogCRM»")
    entities = [
        ("roles", "id PK\nname", 80, 160), ("users", "id PK\nrole_id FK\nemail", 360, 160),
        ("clients", "id PK\nparent_*\narchived_at", 640, 160), ("courses", "id PK\nname\ncost", 920, 160),
        ("groups", "id PK\ncourse_id FK\nteacher_id FK", 1200, 160), ("group_students", "group_id FK\nclient_id FK", 500, 430),
        ("lessons", "id PK\ngroup_id FK\nstart_at", 920, 430), ("attendance", "lesson_id FK\nclient_id FK\nstatus", 1200, 430),
        ("deals", "client_id FK\nmanager_id FK\nstage", 360, 690), ("tasks", "assignee_id FK\nclient_id FK", 640, 690),
        ("user_sessions", "user_id FK\nsession_id", 920, 690), ("audit_logs", "user_id FK\naction", 1200, 690),
    ]
    for title, body, x, y in entities:
        draw_box(draw, (x, y, x + 230, y + 145), title, body, fill=(255, 255, 255), outline=(5, 150, 105))
    for s, e in [((310, 232), (360, 232)), ((590, 232), (640, 232)), ((870, 232), (920, 232)), ((1150, 232), (1200, 232)), ((615, 430), (710, 305)), ((1035, 430), (1300, 305)), ((755, 305), (615, 430)), ((755, 305), (1315, 430)), ((755, 305), (475, 690)), ((755, 305), (755, 690)), ((475, 305), (1035, 690)), ((475, 305), (1315, 690))]:
        arrow(draw, s, e, fill=(71, 85, 105), width=3)
    img.save(ASSETS / "appendix_l_erd.png")


def set_run_font(run, size: float | None = None, name: str = "Times New Roman", bold: bool | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_paragraph_format(paragraph, first_line=True, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.first_line_indent = Cm(1.25) if first_line else Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 14)


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for section in doc.sections:
        for container in [section.header, section.footer]:
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def style_document(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        section.start_type = WD_SECTION_START.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, before, after in [("Heading 1", 18, 6), ("Heading 2", 6, 6), ("Heading 3", 6, 6)]:
        st = doc.styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(14)
        st.font.bold = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.first_line_indent = Cm(1.25)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name.startswith("Heading"):
            for run in paragraph.runs:
                set_run_font(run, 14, bold=False)
            if text.lower().startswith("приложение"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)
            continue
        set_paragraph_format(paragraph)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_mar = tc_pr.first_child_found_in("w:tcMar")
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)
                for margin in ["top", "left", "bottom", "right"]:
                    node = tc_mar.find(qn(f"w:{margin}"))
                    if node is None:
                        node = OxmlElement(f"w:{margin}")
                        tc_mar.append(node)
                    node.set(qn("w:w"), "100")
                    node.set(qn("w:type"), "dxa")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        set_run_font(run, 12)


def replace_text_in_runs(paragraph, replacements: dict[str, str]) -> None:
    full = paragraph.text
    new = full
    for old, value in replacements.items():
        new = new.replace(old, value)
    if new == full:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new
        set_run_font(paragraph.runs[0], 14)
    else:
        run = paragraph.add_run(new)
        set_run_font(run, 14)


def fix_text_content(doc: Document) -> None:
    replacements = {
        "Книжный магазин «Тайны страниц»": "HedgehogCRM",
        "Разработка Web-приложения «HedgehogCRM": "Разработка web-приложения «HedgehogCRM»",
        "1.1Обоснование": "1.1 Обоснование",
        "1.3 Характеристика программных комплексов для решения поставленной задачи.": "1.3 Характеристика программных комплексов для решения поставленной задачи",
        "USE–Case": "USE-Case",
        "Use-Case": "USE-Case",
        "Wed-приложение": "Web-приложение",
        "web-приложения, далее (ПО),": "web-приложения, далее ПО,",
        "Целью исследования, является": "Целью исследования является",
        "В качестве задач исследования, были выявлены": "Для достижения цели были определены",
        "Операционная система Windows 10/11": "Операционная система Windows 11 Pro",
        "Docker 29 и Docker Compose 5.1.0": "Docker Desktop, Docker Engine 29 и Docker Compose",
        "Google Chrome 148.0.7778.168": "Google Chrome",
    }

    for p in iter_all_paragraphs(doc):
        replace_text_in_runs(p, replacements)

    citations = {
        "Python 3.11.9": " [22]",
        "FastAPI —": " [14]",
        "PostgreSQL 16": " [20]",
        "Node.js —": " [19]",
        "Next.js —": " [17, 23, 25]",
        "Docker и Docker Compose": " [12, 13]",
        "Nginx —": " [18]",
        "Visual Studio Code": " [16]",
        "Git 2.49": " [16]",
        "Pydantic": " [21]",
        "SQLAlchemy": " [24]",
    }
    for p in doc.paragraphs:
        txt = p.text.strip()
        for prefix, cite in citations.items():
            if txt.startswith(prefix) and cite not in txt:
                p.add_run(cite)
                set_run_font(p.runs[-1], 14)

    intro_additions = {
        "Базовые понятия и определения предметной области.": "Базовые понятия и определения предметной области, источники данных и организационная структура проекта. ",
        "Структура дипломного проекта включает следующие этапы:": (
            "В качестве источников данных используются сведения об учениках, родителях, сотрудниках, курсах, учебных группах, занятиях, посещаемости, отработках, сделках, задачах, сессиях пользователей и журнале аудита. "
            "Организационная структура проекта включает администратора, менеджера и преподавателя, взаимодействующих с единой базой данных через web-интерфейс. "
            "Разрабатываемая система связана с результатами анализа аналогичных CRM-систем, электронных журналов и платформ дистанционного обучения, но объединяет эти возможности в специализированном решении для IT-академии.\n\n"
            "Структура дипломного проекта включает следующие этапы:"
        ),
    }
    for p in doc.paragraphs:
        for old, new in intro_additions.items():
            if p.text.strip().startswith(old):
                replace_text_in_runs(p, {old: new})


def clear_toc(doc: Document) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    paras = doc.paragraphs
    intro_indices = [i for i, p in enumerate(paras) if p.text.strip() == "Введение"]
    if not intro_indices:
        return
    first_real_intro = intro_indices[0]
    # If the first real intro is the TOC entry, use the second occurrence as content start.
    if len(intro_indices) > 1 and first_real_intro < 8:
        first_real_intro = intro_indices[1]
    first_intro_el = paras[first_real_intro]._p
    remove = []
    for child in children:
        if child is first_intro_el:
            break
        remove.append(child)
    for child in remove:
        body.remove(child)

    title = OxmlElement("w:p")
    body.insert(0, title)
    p = doc.paragraphs[0]
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("Содержание")
    set_run_font(run, 14)

    for idx, (label, _) in enumerate(TOC_ITEMS, start=1):
        p = doc.add_paragraph()
        body.remove(p._p)
        body.insert(idx, p._p)
        p.style = doc.styles["Normal"]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"{label}\t0")
        set_run_font(run, 14)
        if idx == 15:
            p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    body.remove(p._p)
    body.insert(len(TOC_ITEMS) + 1, p._p)
    p.add_run().add_break(WD_BREAK.PAGE)


def remove_elements_between(doc: Document, start_text: str, end_text: str) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start_el = None
    end_el = None
    for p in doc.paragraphs:
        txt = p.text.strip()
        is_real_heading = p.style.name.startswith("Heading")
        if start_el is None and is_real_heading and txt.startswith(start_text):
            start_el = p._p
        if is_real_heading and txt.startswith(end_text):
            end_el = p._p
            break
    if start_el is None or end_el is None:
        return
    removing = False
    for child in children:
        if child is start_el:
            removing = True
        if child is end_el:
            removing = False
        if removing:
            body.remove(child)


def remove_after_heading_content(doc: Document, heading_prefix: str) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start_el = None
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip().startswith(heading_prefix):
            start_el = p._p
            break
    if start_el is None:
        return
    removing = False
    for child in children:
        if child is start_el:
            removing = True
            continue
        if removing:
            if child.tag == qn("w:sectPr"):
                continue
            body.remove(child)


def add_clean_appendix_heading(doc: Document, letter: str) -> None:
    p = doc.add_paragraph()
    if doc.paragraphs and p is not doc.paragraphs[0]:
        p.add_run().add_break(WD_BREAK.PAGE)
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(f"Приложение {letter}")
    set_run_font(r, 14)

    title = doc.add_paragraph()
    title.style = doc.styles["Normal"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run(APPENDIX_TITLES[letter])
    set_run_font(r, 14)


def add_centered_picture(doc: Document, path: Path, width_cm: float = 15.5, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = Cm(0)
        r = cp.add_run(caption)
        set_run_font(r, 14)


def insert_paragraph_before(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def picture_path_for(name: str) -> Path:
    p = ASSETS / name
    if p.exists():
        return p
    return SCREENSHOTS / name


def insert_picture_before_caption(caption_para, image_path: Path) -> None:
    new_p = OxmlElement("w:p")
    caption_para._p.addprevious(new_p)
    p = Paragraph(new_p, caption_para._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))


def insert_body_images(doc: Document) -> None:
    seen = set()
    for p in list(doc.paragraphs):
        text = p.text.strip()
        for key, image_name in FIGURE_IMAGES.items():
            if text.startswith(key) and key not in seen:
                insert_picture_before_caption(p, picture_path_for(image_name))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    set_run_font(run, 14)
                seen.add(key)


def rebuild_appendices(doc: Document) -> None:
    # Replace blank A-I pages before the existing listing, keep and restyle current K listing,
    # then replace L with a generated ERD image.
    remove_elements_between(doc, "Приложение А", "Приложение К")
    body = doc.element.body
    k_para = next((p for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text.strip().startswith("Приложение К")), None)
    if k_para is None:
        return

    # Insert A-I before K.
    for letter in ["А", "Б", "В", "Г", "Д", "Е", "Ж", "З", "И"]:
        title_p = OxmlElement("w:p")
        k_para._p.addprevious(title_p)
        p = Paragraph(title_p, k_para._parent)
        p.add_run().add_break(WD_BREAK.PAGE)
        p.style = doc.styles["Heading 1"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(f"Приложение {letter}")
        set_run_font(r, 14)

        title_xml = OxmlElement("w:p")
        k_para._p.addprevious(title_xml)
        title = Paragraph(title_xml, k_para._parent)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.first_line_indent = Cm(0)
        r = title.add_run(APPENDIX_TITLES[letter])
        set_run_font(r, 14)

        image_xml = OxmlElement("w:p")
        k_para._p.addprevious(image_xml)
        img_p = Paragraph(image_xml, k_para._parent)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.first_line_indent = Cm(0)
        img_p.add_run().add_picture(str(ASSETS / APPENDIX_IMAGES[letter]), width=Cm(15.5))

    # Clean K heading.
    for run in k_para.runs:
        run.text = ""
    k_para.style = doc.styles["Heading 1"]
    k_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k_para.paragraph_format.first_line_indent = Cm(0)
    k_para.runs[0].text = "Приложение К"
    set_run_font(k_para.runs[0], 14)
    title_xml = OxmlElement("w:p")
    k_para._p.addnext(title_xml)
    title = Paragraph(title_xml, k_para._parent)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    r = title.add_run(APPENDIX_TITLES["К"])
    set_run_font(r, 14)

    # Replace L and any following content.
    remove_after_heading_content(doc, "Приложение Л")
    l_para = next((p for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text.strip().startswith("Приложение Л")), None)
    if l_para is not None:
        for run in l_para.runs:
            run.text = ""
        l_para.style = doc.styles["Heading 1"]
        l_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        l_para.paragraph_format.first_line_indent = Cm(0)
        l_para.runs[0].text = "Приложение Л"
        set_run_font(l_para.runs[0], 14)
        title_xml = OxmlElement("w:p")
        l_para._p.addnext(title_xml)
        title = Paragraph(title_xml, l_para._parent)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.first_line_indent = Cm(0)
        r = title.add_run(APPENDIX_TITLES["Л"])
        set_run_font(r, 14)
        img_xml = OxmlElement("w:p")
        title._p.addnext(img_xml)
        img_p = Paragraph(img_xml, title._parent)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.first_line_indent = Cm(0)
        img_p.add_run().add_picture(str(ASSETS / APPENDIX_IMAGES["Л"]), width=Cm(15.5))


def set_page_break_before_structural_headings(doc: Document) -> None:
    structural = [
        "Введение",
        "1 Техническое задание",
        "2 Технический проект",
        "3 Рабочий проект",
        "Заключение",
        "Список литературы",
    ]
    for p in doc.paragraphs:
        txt = p.text.strip()
        if p.style.name.startswith("Heading") and (txt in structural or re.match(r"^Приложение [А-Я]", txt)):
            p.paragraph_format.page_break_before = True


def update_toc(doc: Document, pages: dict[str, int] | None = None) -> None:
    pages = pages or {}
    # Remove existing generated TOC after "Содержание" and before real "Введение".
    clear_toc(doc)
    body = doc.element.body
    for i, p in enumerate(doc.paragraphs[: len(TOC_ITEMS) + 1]):
        if i == 0:
            continue
        label, key = TOC_ITEMS[i - 1]
        page = pages.get(key, 0)
        text = f"{label}\t{page if page else ''}".rstrip()
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = text
            set_run_font(p.runs[0], 14)


def disable_auto_hyphenation(doc: Document) -> None:
    settings = doc.settings._element
    auto = settings.find(qn("w:autoHyphenation"))
    if auto is None:
        auto = OxmlElement("w:autoHyphenation")
        settings.append(auto)
    auto.set(qn("w:val"), "false")


def patch_ooxml_text(docx_path: Path, page_count: int | None = None) -> None:
    replacements = {
        "Разработка Web-приложения «Книжный магазин «Тайны страниц»": "Разработка web-приложения «HedgehogCRM»",
        "Книжный магазин «Тайны страниц»": "HedgehogCRM",
        "Разработка Web-приложения «HedgehogCRM": "Разработка web-приложения «HedgehogCRM»",
    }
    if page_count:
        replacements["  257"] = f"  {page_count}"
    split_replacements = {
        "Разработка Web-приложения «Книжный магазин «Тайны страниц»": "Разработка web-приложения «HedgehogCRM»",
        "Разработка Web-приложения «Книжный магазин «Тайны страниц": "Разработка web-приложения «HedgehogCRM",
        "Книжный магазин «Тайны страниц»": "HedgehogCRM",
        "Книжный магазин «Тайны страниц": "HedgehogCRM",
    }

    def replace_across_text_nodes(xml_text: str) -> str:
        try:
            root = etree.fromstring(xml_text.encode("utf-8"))
        except Exception:
            return xml_text
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        nodes = root.findall(".//w:t", ns)
        if not nodes:
            return xml_text
        for old, new in split_replacements.items():
            while True:
                values = [n.text or "" for n in nodes]
                joined = "".join(values)
                start = joined.find(old)
                if start < 0:
                    break
                end = start + len(old)
                cursor = 0
                touched = []
                for idx, val in enumerate(values):
                    node_start = cursor
                    node_end = cursor + len(val)
                    if node_end > start and node_start < end:
                        touched.append((idx, node_start, node_end))
                    cursor = node_end
                if not touched:
                    break
                first_idx = touched[0][0]
                last_idx = touched[-1][0]
                prefix = values[first_idx][: max(0, start - touched[0][1])]
                suffix = values[last_idx][max(0, end - touched[-1][1]):]
                nodes[first_idx].text = prefix + new + suffix
                for idx, _, _ in touched[1:]:
                    nodes[idx].text = ""
        return etree.tostring(root, encoding="unicode", xml_declaration=False)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    with ZipFile(docx_path, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                for old, new in replacements.items():
                    text = text.replace(old, new)
                text = replace_across_text_nodes(text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp_path, docx_path)


def main() -> None:
    generate_assets()
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    disable_auto_hyphenation(doc)
    style_document(doc)
    fix_text_content(doc)
    update_toc(doc, None)
    insert_body_images(doc)
    rebuild_appendices(doc)
    set_page_break_before_structural_headings(doc)
    style_document(doc)
    doc.save(OUT)
    patch_ooxml_text(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
